import json
import asyncio
import ctypes
import os
import random
import re
import imaplib
import smtplib
import subprocess
import tempfile
import time
import threading
import webbrowser
from email.message import EmailMessage
from email import message_from_bytes
from email.utils import parseaddr
from datetime import datetime, timedelta
from pathlib import Path
from urllib.parse import quote_plus

import musiclibrary
from doc_summarizer import summarize_document
from site_generator import generate_website, parse_website_command

try:
    import status_bridge
except ImportError:
    status_bridge = None


def _status(state, detail=""):
    """Push a state update to the OMIRA orb, if the status bridge is
    available. Never raises — a problem here should never interrupt an
    actual Omira command."""
    if status_bridge is not None:
        try:
            status_bridge.set_status(state, detail)
        except Exception:
            pass


# ---------------------------------------------------------------------------
# QUIET TERMINAL
#
# By default Omira prints nothing to the console — startup, listening,
# recognized text, and every internal diagnostic ("[Gemini] ...", "[AI] ...")
# are all routed through log() instead of the builtin print(). The console
# stays silent; the only thing you hear is Omira speaking. Set
# OMIRA_VERBOSE=1 in .env (or as an environment variable) any time you want
# the old diagnostic output back, e.g. while troubleshooting.
# ---------------------------------------------------------------------------
_print = print
VERBOSE = os.getenv("OMIRA_VERBOSE", "0") == "1"


def log(*args, **kwargs):
    if VERBOSE:
        _print(*args, **kwargs)


try:
    from dotenv import load_dotenv
    load_dotenv()  # reads .env in the current folder into os.environ, BEFORE
                    # any os.getenv() calls below run. Without this, every
                    # setting in .env is silently ignored.
except ImportError:
    log("python-dotenv is not installed — .env will not be loaded. Run: pip install python-dotenv")

import pyttsx3
import pyautogui
import pywhatkit
import requests
import speech_recognition as sr

try:
    from openai import OpenAI
except ImportError:
    OpenAI = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import screen_brightness_control as sbc
except ImportError:
    sbc = None

try:
    import edge_tts
except ImportError:
    edge_tts = None

try:
    import playsound
except ImportError:
    playsound = None

try:
    from comtypes import CLSCTX_ALL
    from pycaw.pycaw import AudioUtilities, IAudioEndpointVolume
    PYCAW_AVAILABLE = True
except ImportError:
    PYCAW_AVAILABLE = False

try:
    from vosk import Model, KaldiRecognizer
except ImportError:
    Model = None
    KaldiRecognizer = None

try:
    from faster_whisper import WhisperModel
except ImportError:
    WhisperModel = None



def create_tts_engine():
    tts_engine = pyttsx3.init("sapi5")
    tts_engine.setProperty("rate", int(os.getenv("OMIRA_SPEECH_RATE", "220")))
    tts_engine.setProperty("volume", 1.0)

    for voice in tts_engine.getProperty("voices"):
        voice_id = getattr(voice, "id", "")
        if "ZIRA" in voice_id.upper() or "DAVID" in voice_id.upper():
            tts_engine.setProperty("voice", voice.id)
            break

    return tts_engine


engine = create_tts_engine()

OLLAMA_URL = os.getenv("OMIRA_OLLAMA_URL", "http://localhost:11434/api/chat")
OLLAMA_MODEL = os.getenv("OMIRA_OLLAMA_MODEL", "llama3.1")

# AI cascade, tried in this order: Gemini -> OpenAI -> local Ollama. Each is
# skipped automatically if its key/server isn't configured, so you only need
# ONE of the three working for questions/documents/planning to work at all.
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")

# Google renames/retires Gemini model IDs often enough that pinning to a
# single one risks a 404 ("model not found") a few months later — which
# looks exactly like "Omira stopped answering questions". So there is
# always a fallback list. If you set OMIRA_GEMINI_MODEL yourself, that
# model is tried FIRST (so you still get your preferred/cheaper model when
# it works), but the standard candidates are tried after it instead of
# giving up — a single bad pin (e.g. an unstable "-latest" alias that got
# retired) no longer kills Gemini entirely.
_env_gemini_model = os.getenv("OMIRA_GEMINI_MODEL", "")
_DEFAULT_GEMINI_CANDIDATES = ["gemini-2.5-flash", "gemini-2.0-flash", "gemini-flash-latest", "gemini-1.5-flash"]
if _env_gemini_model:
    GEMINI_MODEL_CANDIDATES = [_env_gemini_model] + [
        m for m in _DEFAULT_GEMINI_CANDIDATES if m != _env_gemini_model
    ]
else:
    GEMINI_MODEL_CANDIDATES = list(_DEFAULT_GEMINI_CANDIDATES)
_gemini_working_model = None  # set the first time a candidate succeeds

# After a 429 (quota/rate-limit) response, stop calling Gemini for this many
# seconds instead of retrying it on every single command. Without this,
# hands-free mode (no wake word) keeps re-hitting an already-exhausted quota
# on every stray phrase the mic picks up, which is the single biggest cause
# of "Gemini always 429s" — it also means Omira fails over to OpenAI/Ollama
# instantly instead of waiting on a doomed request first.
GEMINI_COOLDOWN_SECONDS = int(os.getenv("OMIRA_GEMINI_COOLDOWN_SECONDS", "90"))
_gemini_cooldown_until = 0.0

# Set whenever a backend fails in a way that has a clear, fixable cause
# (bad/expired key, quota, etc.), so the spoken/failure message can tell you
# WHAT to fix instead of just "try again" — no need to turn on
# OMIRA_VERBOSE=1 just to find out why a question went unanswered.
_LAST_AI_ERROR = ""

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
OPENAI_MODEL = os.getenv("OMIRA_OPENAI_MODEL", "gpt-4o-mini")
openai_client = OpenAI(api_key=OPENAI_API_KEY) if (OpenAI and OPENAI_API_KEY) else None

OMIRA_PERSONA = (
    "You are Omira, a calm, precise personal assistant in the style of "
    "Tony Stark's AI. Address the user as 'sir'. Be brief, confident, and "
    "a little dry — never chatty, never apologetic, never verbose."
)
EMAIL_ADDRESS = os.getenv("OMIRA_EMAIL_ADDRESS", "")
EMAIL_PASSWORD = os.getenv("OMIRA_EMAIL_PASSWORD", "")
SMTP_SERVER = os.getenv("OMIRA_SMTP_SERVER", "smtp.gmail.com")
SMTP_PORT = int(os.getenv("OMIRA_SMTP_PORT", "587"))
IMAP_SERVER = os.getenv("OMIRA_IMAP_SERVER", "imap.gmail.com")
IMAP_PORT = int(os.getenv("OMIRA_IMAP_PORT", "993"))

# Set by check_email()/read_email_from() so "reply saying ..." knows who to
# reply to without the user repeating the sender's address.
LAST_READ_EMAIL = {"sender": None, "subject": None}

APP_COMMANDS = {
    "calculator": ["calc.exe"],
    "notepad": ["notepad.exe"],
    "paint": ["mspaint.exe"],
    "command prompt": ["cmd.exe"],
    "terminal": ["powershell.exe"],
    "file explorer": ["explorer.exe"],
    "settings": ["cmd", "/c", "start", "", "ms-settings:"],
    "control panel": ["control.exe"],
    "task manager": ["taskmgr.exe"],
    "device manager": ["cmd", "/c", "start", "", "devmgmt.msc"],
    "this pc": ["explorer.exe", "shell:MyComputerFolder"],
    # Specific Settings pages — "full laptop control" in practice means
    # being able to jump straight to the page you actually want, not just
    # the Settings home screen.
    "wifi settings": ["cmd", "/c", "start", "", "ms-settings:network-wifi"],
    "bluetooth settings": ["cmd", "/c", "start", "", "ms-settings:bluetooth"],
    "display settings": ["cmd", "/c", "start", "", "ms-settings:display"],
    "sound settings": ["cmd", "/c", "start", "", "ms-settings:sound"],
    "battery settings": ["cmd", "/c", "start", "", "ms-settings:batterysaver"],
    "storage settings": ["cmd", "/c", "start", "", "ms-settings:storagesense"],
    "update settings": ["cmd", "/c", "start", "", "ms-settings:windowsupdate"],
    "apps settings": ["cmd", "/c", "start", "", "ms-settings:appsfeatures"],
    "network settings": ["cmd", "/c", "start", "", "ms-settings:network-status"],
    "printer settings": ["cmd", "/c", "start", "", "ms-settings:printers"],
}

PENDING_ACTION = None
REMINDERS = []
REMINDER_LOCK = threading.Lock()
SPEAK_LOCK = threading.Lock()
SPEECH_BACKEND = os.getenv("OMIRA_SPEECH_BACKEND", "edge").lower()
FAST_MODE = os.getenv("OMIRA_FAST_MODE", "1") == "1"

# ---------------------------------------------------------------------------
# LANGUAGE — English / Hindi / Marathi
#
# Two separate things, because they need different solutions:
#   1. LISTENING: Google's free speech API needs ONE language code told to it
#      in advance for each recording — it can't auto-detect which of three
#      languages you're about to speak. So CURRENT_LANGUAGE is a persistent
#      setting you switch with a voice command ("speak in hindi").
#   2. REPLYING: once your words are text, Omira reads the SCRIPT (Latin vs
#      Devanagari) to decide what language to answer in, automatically, no
#      command needed. Devanagari can't distinguish Hindi from Marathi by
#      script alone, so that case falls back to CURRENT_LANGUAGE.
# ---------------------------------------------------------------------------

LANGUAGE_STT_CODES = {"english": "en-IN", "hindi": "hi-IN", "marathi": "mr-IN"}

# Single voice culture used for EVERY language, on every TTS backend (edge,
# PowerShell, and the pyttsx3 last-resort fallback). Previously the
# PowerShell and pyttsx3 fallbacks each picked a different installed voice
# per language even after edge-tts was pinned to one voice — so if edge-tts
# wasn't actually installed (it was missing from requirements.txt), Omira
# silently fell back to a path that still swapped voices. Fixed on both
# counts below.
OMIRA_VOICE_CULTURE = os.getenv("OMIRA_VOICE_CULTURE", "hi-IN")

# ONE real Indian-accented neural voice via edge-tts (free, no API key —
# uses Microsoft Edge's online TTS service), used for English, Hindi, and
# Marathi alike, so Omira doesn't switch voices mid-conversation when the
# language changes. This is what actually fixes the "sounds British"
# complaint — Windows' built-in SAPI5 voices are US/UK English by default
# and only sound Indian if you've separately installed a Hindi language
# pack, which most people haven't.
#
# hi-IN-SwaraNeural is the default because it's Microsoft's "multilingual"
# Indian voice — it's built to read English words cleanly inside Hindi
# speech (rather than mangling them phonetically), and since Marathi shares
# Devanagari script and most sounds with Hindi, it reads Marathi clearly
# too. Override with OMIRA_EDGE_VOICE in .env if you'd rather use a
# different single voice — male alternative: hi-IN-MadhurNeural.
OMIRA_EDGE_VOICE = os.getenv("OMIRA_EDGE_VOICE", "hi-IN-SwaraNeural")
EDGE_VOICES = {
    "english": OMIRA_EDGE_VOICE,
    "hindi": OMIRA_EDGE_VOICE,
    "marathi": OMIRA_EDGE_VOICE,
}
LANGUAGE_PROMPT_NAMES = {"english": "English", "hindi": "Hindi", "marathi": "Marathi"}

CURRENT_LANGUAGE = os.getenv("OMIRA_DEFAULT_LANGUAGE", "english").lower()
if CURRENT_LANGUAGE not in LANGUAGE_STT_CODES:
    CURRENT_LANGUAGE = "english"

# Devanagari script is a sure sign of Hindi/Marathi, but a lot of people
# (like "gravity kya hai") type or speak it phonetically in Latin letters
# instead. These common function words catch that case. Marathi's are more
# distinctive (kay/ahe/kuthe vs Hindi's kya/hai/kahan), so Marathi is
# checked first to reduce misclassifying Marathi as Hindi.
MARATHI_MARKER_WORDS = {
    "kay", "ahe", "nahi", "kasa", "kashi", "kase", "kon", "kuthe", "kadhi",
    "mala", "tula", "amhi", "tumhi", "amhala", "tumhala", "zala", "zali",
    "zale", "changla", "kiti", "tuza", "maza",
}
HINDI_MARKER_WORDS = {
    "kya", "hai", "kaise", "kaisi", "kaisa", "kaun", "kab", "kahan", "kyun",
    "kyu", "hoon", "hun", "tha", "thi", "the", "aap", "tum", "mera", "meri",
    "tera", "teri", "uska", "uski", "hum", "nahin", "haan", "matlab",
    "acha", "accha", "theek", "batao", "bata", "karo", "raha", "rahi",
    "rahe", "mujhe", "tumhe", "kripya", "kitna", "kitni",
}


def detect_reply_language(text):
    """Devanagari script -> Hindi, unless Marathi is the active language
    (script alone can't tell Hindi and Marathi apart). Romanized Hindi/
    Marathi (Latin letters, e.g. 'gravity kya hai') is caught via common
    function words. Otherwise -> English."""
    if re.search(r"[\u0900-\u097F]", text):
        return "marathi" if CURRENT_LANGUAGE == "marathi" else "hindi"

    words = set(re.findall(r"[a-zA-Z']+", text.lower()))
    if words & MARATHI_MARKER_WORDS:
        return "marathi"
    if words & HINDI_MARKER_WORDS:
        return "hindi"
    return "english"


def parse_language_switch_command(command):
    lowered = command.lower().strip()
    if re.search(r"\b(speak|talk|reply|respond)\b.*\bhindi\b", lowered) or "hindi mein baat" in lowered:
        return "hindi"
    if re.search(r"\b(speak|talk|reply|respond)\b.*\bmarathi\b", lowered) or "marathi madhe" in lowered:
        return "marathi"
    if re.search(r"\b(speak|talk|reply|respond)\b.*\benglish\b", lowered) or "english mein baat" in lowered:
        return "english"
    return None


STT_BACKEND = os.getenv("OMIRA_STT_BACKEND", "google").lower()
VOSK_MODEL_PATH = os.getenv("OMIRA_VOSK_MODEL", "")
WHISPER_MODEL = os.getenv("OMIRA_WHISPER_MODEL", "base.en")


CONTACTS_FILE = Path(__file__).with_name("contacts.json")
CONTACTS = {}
REMINDERS_FILE = Path(__file__).with_name("reminders.json")

# Seconds pywhatkit waits for web.whatsapp.com to load before it types and
# sends the message. Slow connections may need this bumped up.
WHATSAPP_WAIT_SECONDS = int(os.getenv("OMIRA_WHATSAPP_WAIT_SECONDS", "20"))

REMINDER_TIME_FORMATS = [
    "%I:%M %p",
    "%I %p",
    "%H:%M",
    "%H:%M:%S",
]

pyautogui.FAILSAFE = True

VK_VOLUME_MUTE = 0xAD
VK_VOLUME_DOWN = 0xAE
VK_VOLUME_UP = 0xAF
VK_MEDIA_PLAY_PAUSE = 0xB3


def speak(text, language=None):
    global engine

    message = str(text).strip()
    if not message:
        return

    language = language or CURRENT_LANGUAGE
    # Always shown — this is the conversation transcript, not a diagnostic.
    _print(f"Omira: {message}")
    _status("responding", message)

    if SPEECH_BACKEND == "edge":
        if speak_with_edge_tts(message, language):
            return
        log("edge-tts failed (needs internet + `pip install edge-tts playsound==1.2.2`), trying PowerShell fallback.")

    if SPEECH_BACKEND in ("edge", "powershell"):
        if speak_with_powershell(message, language):
            return
        log("PowerShell speech failed, trying pyttsx3 fallback.")

    set_pyttsx3_voice_for_language(language)
    for _ in range(2):
        try:
            with SPEAK_LOCK:
                engine.say(message)
                engine.runAndWait()
            return
        except Exception as exc:
            log(f"TTS error: {exc}")
            try:
                engine = create_tts_engine()
            except Exception as init_exc:
                log(f"TTS reinitialize failed: {init_exc}")

    if not speak_with_powershell(message, language):
        log("Speech failed on all backends.")


def speak_async(text, language=None):
    threading.Thread(target=speak, args=(text, language), daemon=True).start()


async def _edge_tts_save(text, voice, path):
    communicate = edge_tts.Communicate(text, voice)
    await communicate.save(path)


def speak_with_edge_tts(text, language="english"):
    """Real Indian-accented neural voice via Microsoft's free online TTS.
    Needs internet and `pip install edge-tts playsound==1.2.2`."""
    if edge_tts is None or playsound is None:
        return False

    voice = EDGE_VOICES.get(language, EDGE_VOICES["english"])
    audio_path = os.path.join(tempfile.gettempdir(), "omira_speech.mp3")

    try:
        with SPEAK_LOCK:
            asyncio.run(_edge_tts_save(text, voice, audio_path))
            playsound.playsound(audio_path)
        return True
    except Exception as exc:
        log(f"edge-tts request failed: {exc}")
        return False


_pyttsx3_voice_locked = False


def set_pyttsx3_voice_for_language(language):
    """Last-resort fallback voice (Windows SAPI5, offline). Picks ONE voice
    the first time this runs and keeps it for every language after that —
    it deliberately ignores `language` so it can't reintroduce the
    multiple-voices problem edge-tts/PowerShell are meant to avoid. Prefers
    a voice whose name matches OMIRA_VOICE_CULTURE (e.g. an installed Hindi
    voice) if one exists, otherwise keeps whatever create_tts_engine() set."""
    global _pyttsx3_voice_locked
    if _pyttsx3_voice_locked:
        return
    try:
        target = OMIRA_VOICE_CULTURE.split("-")[0].upper()  # e.g. "HI"
        for voice in engine.getProperty("voices"):
            name = (getattr(voice, "name", "") or "").upper()
            voice_id = (getattr(voice, "id", "") or "").upper()
            if target in name or target in voice_id:
                engine.setProperty("voice", voice.id)
                break
    except Exception:
        pass
    finally:
        _pyttsx3_voice_locked = True


def speak_with_powershell(text, language="english"):
    # Wrapped in SPEAK_LOCK (same lock edge-tts and pyttsx3 use) so this
    # can never overlap with another speak() call running on a different
    # thread — e.g. speak_async("Opening YouTube") racing a normal reply.
    # Without this lock, two PowerShell speech processes could run at once,
    # which is what caused hearing two different sentences simultaneously.
    with SPEAK_LOCK:
        # Always the same voice culture, regardless of language — see
        # OMIRA_VOICE_CULTURE above. `language` param kept for call-site
        # compatibility with speak(), but intentionally unused for voice choice.
        escaped = text.replace("'", "''")
        culture_prefix = OMIRA_VOICE_CULTURE
        ps_script = (
            "Add-Type -AssemblyName System.Speech; "
            "$speaker = New-Object System.Speech.Synthesis.SpeechSynthesizer; "
            f"$voice = $speaker.GetInstalledVoices() | Where-Object "
            f"{{ $_.VoiceInfo.Culture.Name -like '{culture_prefix}*' }} | Select-Object -First 1; "
            "if ($voice) { $speaker.SelectVoice($voice.VoiceInfo.Name) } "
            f"$speaker.Speak('{escaped}')"
        )
        command = ["powershell", "-NoProfile", "-Command", ps_script]

        try:
            result = subprocess.run(command, check=False, capture_output=True, text=True)
            if result.returncode == 0:
                return True

            log(f"PowerShell speech failed with return code {result.returncode}: {result.stderr.strip()}")
            return False
        except Exception as exc:
            log(f"PowerShell speech fallback failed: {exc}")
            return False



# ---------------------------------------------------------------------------
# OPTIONAL LOCAL SPEECH RECOGNITION
# Set OMIRA_STT_BACKEND=vosk or faster-whisper after installing the backend.
# Google remains the default so the script works with the existing setup.
# ---------------------------------------------------------------------------

LOCAL_VOSK_MODEL = None
LOCAL_WHISPER_MODEL = None


def initialize_local_stt():
    global LOCAL_VOSK_MODEL, LOCAL_WHISPER_MODEL

    if STT_BACKEND == "vosk":
        if Model is None:
            log("Vosk is not installed; falling back to Google STT.")
            return
        if not VOSK_MODEL_PATH:
            log("OMIRA_VOSK_MODEL is not set; falling back to Google STT.")
            return
        try:
            LOCAL_VOSK_MODEL = Model(VOSK_MODEL_PATH)
            log("Local Vosk speech recognition ready.")
        except Exception as exc:
            log(f"Vosk initialization failed: {exc}")
            LOCAL_VOSK_MODEL = None

    elif STT_BACKEND in {"whisper", "faster-whisper"}:
        if WhisperModel is None:
            log("faster-whisper is not installed; falling back to Google STT.")
            return
        try:
            LOCAL_WHISPER_MODEL = WhisperModel(
                WHISPER_MODEL,
                device="cpu",
                compute_type="int8",
            )
            log(f"Local Whisper model '{WHISPER_MODEL}' ready.")
        except Exception as exc:
            log(f"Whisper initialization failed: {exc}")
            LOCAL_WHISPER_MODEL = None


def recognize_audio(audio, recognizer):
    """Recognize audio using the configured backend."""
    if STT_BACKEND == "vosk" and LOCAL_VOSK_MODEL is not None:
        import wave
        import json as _json
        wav_bytes = audio.get_wav_data()
        import io
        with wave.open(io.BytesIO(wav_bytes), "rb") as wf:
            rec = KaldiRecognizer(LOCAL_VOSK_MODEL, wf.getframerate())
            rec.SetWords(False)
            while True:
                data = wf.readframes(4000)
                if not data:
                    break
                rec.AcceptWaveform(data)
            result = _json.loads(rec.FinalResult())
            return result.get("text", "").strip()

    if STT_BACKEND in {"whisper", "faster-whisper"} and LOCAL_WHISPER_MODEL is not None:
        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as temp:
            temp.write(audio.get_wav_data())
            temp_path = temp.name
        try:
            segments, _ = LOCAL_WHISPER_MODEL.transcribe(
                temp_path,
                beam_size=1,
                best_of=1,
                temperature=0,
                vad_filter=True,
            )
            return " ".join(segment.text.strip() for segment in segments).strip()
        finally:
            try:
                os.remove(temp_path)
            except OSError:
                pass

    return recognizer.recognize_google(audio, language=LANGUAGE_STT_CODES.get(CURRENT_LANGUAGE, "en-IN"))


def create_recognizer():
    recognizer = sr.Recognizer()
    recognizer.dynamic_energy_threshold = True
    # 0.35s was cutting the recording off on the natural micro-pause inside
    # a question ("what is ... AI") before the last word was even spoken,
    # so short commands worked but questions got truncated. 0.7s still feels
    # snappy for single commands but gives a full question room to breathe.
    # Tune with OMIRA_PAUSE_THRESHOLD if it still feels off in either direction.
    recognizer.pause_threshold = float(os.getenv("OMIRA_PAUSE_THRESHOLD", "0.7"))
    recognizer.non_speaking_duration = min(0.4, recognizer.pause_threshold)
    recognizer.operation_timeout = 8
    return recognizer


RECOGNIZER = None  # set at startup in __main__


def listen_once(timeout=4, phrase_time_limit=6):
    """Listen for a single short reply (used for the yes/no follow-up prompt).
    Returns '' on silence/timeout instead of raising."""
    global RECOGNIZER
    if RECOGNIZER is None:
        RECOGNIZER = create_recognizer()
    try:
        with sr.Microphone() as source:
            audio = RECOGNIZER.listen(source, timeout=timeout, phrase_time_limit=phrase_time_limit)
        return recognize_audio(audio, RECOGNIZER).strip()
    except Exception:
        return ""


def load_contacts():
    if CONTACTS_FILE.exists():
        try:
            with CONTACTS_FILE.open("r", encoding="utf-8") as file_handle:
                contacts = json.load(file_handle)
                if isinstance(contacts, dict):
                    return contacts
        except Exception:
            pass
    return {}


def load_reminders():
    if not REMINDERS_FILE.exists():
        return []

    try:
        with REMINDERS_FILE.open("r", encoding="utf-8") as file_handle:
            stored_reminders = json.load(file_handle)
    except Exception:
        return []

    reminders = []
    if not isinstance(stored_reminders, list):
        return reminders

    for item in stored_reminders:
        if not isinstance(item, dict):
            continue

        task = item.get("task")
        time_text = item.get("time")
        if not task or not time_text:
            continue

        try:
            reminder_time = datetime.fromisoformat(time_text)
        except ValueError:
            continue

        reminders.append({"task": task, "time": reminder_time})

    return reminders


def save_reminders():
    with REMINDER_LOCK:
        serializable_reminders = [
            {"task": reminder["task"], "time": reminder["time"].isoformat()}
            for reminder in REMINDERS
        ]

    try:
        with REMINDERS_FILE.open("w", encoding="utf-8") as file_handle:
            json.dump(serializable_reminders, file_handle, indent=2)
    except Exception:
        pass


def get_battery_status():
    class SYSTEM_POWER_STATUS(ctypes.Structure):
        _fields_ = [
            ("ACLineStatus", ctypes.c_ubyte),
            ("BatteryFlag", ctypes.c_ubyte),
            ("BatteryLifePercent", ctypes.c_ubyte),
            ("SystemStatusFlag", ctypes.c_ubyte),
            ("BatteryLifeTime", ctypes.c_ulong),
            ("BatteryFullLifeTime", ctypes.c_ulong),
        ]

    status = SYSTEM_POWER_STATUS()
    if not ctypes.windll.kernel32.GetSystemPowerStatus(ctypes.byref(status)):
        return "I could not read the battery status."

    if status.BatteryLifePercent == 255:
        percent_text = "unknown battery level"
    else:
        percent_text = f"{status.BatteryLifePercent} percent battery"

    if status.ACLineStatus == 1:
        power_text = "and it is charging"
    elif status.ACLineStatus == 0:
        power_text = "and it is running on battery"
    else:
        power_text = ""

    return f"Your laptop battery is at {percent_text} {power_text}".strip()


def get_system_info():
    """Actual laptop specs — OS, CPU, RAM, disk. No AI call involved, so
    this always works even if OpenAI/Ollama are unreachable."""
    import platform
    import shutil

    os_text = f"{platform.system()} {platform.release()}"
    cpu_text = platform.processor() or "an unknown processor"

    ram_text = "unknown RAM"
    try:
        class MEMORYSTATUSEX(ctypes.Structure):
            _fields_ = [
                ("dwLength", ctypes.c_ulong),
                ("dwMemoryLoad", ctypes.c_ulong),
                ("ullTotalPhys", ctypes.c_ulonglong),
                ("ullAvailPhys", ctypes.c_ulonglong),
                ("ullTotalPageFile", ctypes.c_ulonglong),
                ("ullAvailPageFile", ctypes.c_ulonglong),
                ("ullTotalVirtual", ctypes.c_ulonglong),
                ("ullAvailVirtual", ctypes.c_ulonglong),
                ("sullAvailExtendedVirtual", ctypes.c_ulonglong),
            ]

        mem_status = MEMORYSTATUSEX()
        mem_status.dwLength = ctypes.sizeof(MEMORYSTATUSEX)
        if ctypes.windll.kernel32.GlobalMemoryStatusEx(ctypes.byref(mem_status)):
            total_gb = mem_status.ullTotalPhys / (1024 ** 3)
            used_percent = mem_status.dwMemoryLoad
            ram_text = f"{total_gb:.1f} gigabytes of RAM, {used_percent} percent in use"
    except Exception:
        pass

    disk_text = "unknown disk space"
    try:
        total, used, free = shutil.disk_usage("C:\\" if os.name == "nt" else "/")
        disk_text = f"{free / (1024 ** 3):.0f} gigabytes free out of {total / (1024 ** 3):.0f}"
    except Exception:
        pass

    return (
        f"You are running {os_text}, sir, on a {cpu_text} processor, "
        f"with {ram_text}, and {disk_text} of disk space free."
    )


def get_weather(place=None):
    location = quote_plus(place.strip()) if place else ""
    url = f"https://wttr.in/{location}?format=j1" if location else "https://wttr.in/?format=j1"
    response = requests.get(url, timeout=15)
    response.raise_for_status()
    payload = response.json()

    current = payload["current_condition"][0]
    forecast = payload["weather"][0]
    description = current["weatherDesc"][0]["value"]
    temp_c = current["temp_C"]
    feels_like = current["FeelsLikeC"]
    max_temp = forecast["maxtempC"]
    min_temp = forecast["mintempC"]

    if place:
        prefix = f"The weather in {place}"
    else:
        prefix = "Today's weather"

    return (
        f"{prefix} is {description}, {temp_c} degrees Celsius, feels like {feels_like}, "
        f"with a high of {max_temp} and a low of {min_temp}."
    )


def parse_weather_command(command):
    match = re.search(r"\bweather(?: in| of)?\s+(?P<place>.+)$", command, flags=re.IGNORECASE)
    if match:
        place = match.group("place").strip()
        if place:
            return place

    if re.search(r"\bwhat(?:'s| is)? the weather\b", command, flags=re.IGNORECASE):
        return ""

    if re.fullmatch(r"weather", command.strip(), flags=re.IGNORECASE):
        return ""

    return None


def speak_weather(place=None):
    try:
        speak(get_weather(place))
    except Exception:
        if place:
            speak(f"I could not get the weather for {place} right now.")
        else:
            speak("I could not get today's weather right now.")


def get_reminder_datetime(time_text):
    cleaned_time = time_text.strip().lower().replace(".", "")

    for time_format in REMINDER_TIME_FORMATS:
        try:
            parsed_time = datetime.strptime(cleaned_time, time_format).time()
            reminder_datetime = datetime.combine(datetime.now().date(), parsed_time)
            if reminder_datetime <= datetime.now():
                reminder_datetime += timedelta(days=1)
            return reminder_datetime
        except ValueError:
            continue

    return None


def schedule_reminder(task, reminder_datetime):
    reminder = {"task": task.strip(), "time": reminder_datetime}
    with REMINDER_LOCK:
        REMINDERS.append(reminder)
    save_reminders()
    speak(f"Reminder set for {task} at {reminder_datetime.strftime('%I:%M %p')}")


def reminder_worker():
    while True:
        now = datetime.now()
        due_reminders = []

        with REMINDER_LOCK:
            pending_reminders = []
            for reminder in REMINDERS:
                if reminder["time"] <= now:
                    due_reminders.append(reminder)
                else:
                    pending_reminders.append(reminder)
            REMINDERS[:] = pending_reminders

        for reminder in due_reminders:
            speak(f"Reminder: {reminder['task']}")

        if due_reminders:
            save_reminders()

        time.sleep(1)


def parse_reminder_command(command):
    patterns = [
        r"(?:set\s+)?remind(?:er| me)?\s+to\s+(?P<task>.+?)\s+in\s+(?P<amount>\d+)\s+(?P<unit>seconds?|minutes?|hours?)$",
        r"(?:set\s+)?remind(?:er| me)?\s+to\s+(?P<task>.+?)\s+at\s+(?P<time>.+)$",
        r"(?:set\s+)?remind(?:er| me)?\s+(?P<task>.+?)\s+at\s+(?P<time>.+)$",
    ]

    for pattern in patterns:
        match = re.search(pattern, command, flags=re.IGNORECASE)
        if not match:
            continue

        task = match.group("task").strip()
        amount = match.groupdict().get("amount")
        unit = match.groupdict().get("unit")
        time_text = match.groupdict().get("time")

        if amount and unit:
            amount_value = int(amount)
            if unit.startswith("second"):
                reminder_datetime = datetime.now() + timedelta(seconds=amount_value)
            elif unit.startswith("minute"):
                reminder_datetime = datetime.now() + timedelta(minutes=amount_value)
            else:
                reminder_datetime = datetime.now() + timedelta(hours=amount_value)
            return task, reminder_datetime

        if time_text:
            reminder_datetime = get_reminder_datetime(time_text)
            if reminder_datetime:
                return task, reminder_datetime

    return None


def parse_message_command(command):
    patterns = [
        r"send\s+(?P<message>.+?)\s+message\s+to\s+(?P<person>.+?)\s+(?:on|via)\s+whatsapp$",
        r"send\s+message\s+to\s+(?P<person>.+?)\s+(?:on|via)\s+whatsapp\s+(?:saying|that says)\s+(?P<message>.+)$",
        r"(?:send\s+)?(?:whatsapp\s+)?message\s+to\s+(?P<person>.+?)\s+(?:saying|that says|message)\s+(?P<message>.+)$",
        r"(?:send\s+)?(?:whatsapp\s+)?message\s+to\s+(?P<person>.+?)\s*[:,-]\s*(?P<message>.+)$",
        r"(?:send\s+)?(?:whatsapp\s+)?message\s+to\s+(?P<person>.+?)\s+(?:on|via)\s+whatsapp$",
        r"(?:text|message)\s+(?P<person>.+?)\s+(?:saying|that says)\s+(?P<message>.+)$",
        r"send\s+(?P<message>.+?)\s+to\s+(?P<person>.+?)\s+on\s+whatsapp$",
        r"send\s+(?P<message>.+?)\s+to\s+(?P<person>.+?)\s+via\s+whatsapp$",
    ]

    for pattern in patterns:
        match = re.search(pattern, command, flags=re.IGNORECASE)
        if match:
            person = match.group("person").strip()
            message = match.groupdict().get("message", "").strip()

            if person.lower().endswith(" on whatsapp"):
                person = person[:-12].strip()

            if message.startswith('"') and message.endswith('"') and len(message) >= 2:
                message = message[1:-1].strip()

            if message.startswith("'") and message.endswith("'") and len(message) >= 2:
                message = message[1:-1].strip()

            if not message:
                message = "Hi"

            return person, message

    return None


def resolve_contact(person):
    cleaned_person = person.strip()
    normalized_phone = re.sub(r"[^\d+]", "", cleaned_person)
    if normalized_phone.isdigit() or normalized_phone.startswith("+"):
        return normalized_phone

    for name, phone in CONTACTS.items():
        if name.lower() == cleaned_person.lower().strip():
            return phone

    return None


# ---------------------------------------------------------------------------
# WHATSAPP — SEND BY NAME (via contacts.json)
#
# Unlike the old search-and-send approach (pixel-coordinate clicking,
# removed from this build), this only needs a phone number, which it gets
# from contacts.json via resolve_contact(). It hands off to pywhatkit,
# which opens web.whatsapp.com in your default browser, waits for it to
# load, types the message into the already-open chat, and hits Enter.
#
# Requirements: default browser already logged into WhatsApp Web, and the
# person's name spelled in contacts.json exactly as you'll say it (matching
# is case-insensitive, but not fuzzy).
# ---------------------------------------------------------------------------

def send_whatsapp_message(phone, message, person=None):
    label = person or phone
    try:
        pywhatkit.sendwhatmsg_instantly(
            phone_no=phone,
            message=message,
            wait_time=WHATSAPP_WAIT_SECONDS,
            tab_close=True,
            close_time=3,
        )
        speak(f"Message sent to {label} on WhatsApp, sir.")
    except Exception as exc:
        log(f"[WhatsApp] send to {label} failed: {exc}")
        speak(f"I couldn't send that WhatsApp message to {label}, sir.")


def handle_startup_briefing():
    speak(get_battery_status())

    try:
        speak(get_weather())
    except Exception:
        speak("I could not get today's weather right now.")

    speak("What do you want to do today?")


def speak_daily_briefing():
    speak(get_battery_status())
    speak_weather()


def test_voice_output():
    speak("Hello sir. My voice is working.")


def open_website(url):
    webbrowser.open(url)


def open_app(app_name):
    command = APP_COMMANDS.get(app_name.lower().strip())
    if not command:
        speak(f"I do not know how to open {app_name} yet.")
        return

    subprocess.Popen(command)
    speak(f"Opening {app_name}")


def open_file(target_path):
    path = Path(target_path).expanduser()
    if not path.exists():
        speak(f"I could not find {target_path}.")
        return

    os.startfile(str(path))
    speak(f"Opening {path.name}")


RANDOM_SONG_PHRASES = {
    "any song", "a song", "random song", "some song", "some music",
    "music", "anything", "something", "a random song",
}

# Trailing words people naturally add that aren't part of the song title.
_SONG_FILLER_SUFFIXES = (
    " on youtube", " on you tube", " song", " songs", " track", " music", " please",
)


def normalize_song_query(text):
    """Strip filler like 'song'/'on youtube' so 'Shape Of You song on YouTube'
    matches the library key 'shape of you'. Stops early if what's left is
    already a generic 'play something' phrase like 'any song'."""
    query = text.lower().strip()
    while True:
        if query in RANDOM_SONG_PHRASES:
            return query
        stripped = False
        for suffix in _SONG_FILLER_SUFFIXES:
            if query.endswith(suffix):
                query = query[: -len(suffix)].strip()
                stripped = True
                break  # re-check membership before stripping further
        if not stripped:
            return query


def play_song(requested_text):
    query = normalize_song_query(requested_text)

    if not query or query in RANDOM_SONG_PHRASES:
        play_random_song()
        return

    link = musiclibrary.music.get(query)

    # No exact key match — try a loose match before giving up (e.g. "arijit
    # singh songs" -> library key "arijit singh").
    if not link:
        for key, url in musiclibrary.music.items():
            if query in key or key in query:
                link = url
                query = key
                break

    if link:
        speak(f"Playing {query}")
        webbrowser.open(link)
        return

    # Genuinely not in the library — search YouTube directly instead of
    # just refusing, so a spoken request always plays *something*.
    speak(f"That is not in my library, sir. Searching YouTube for {requested_text}.")
    play_youtube(requested_text)


def play_youtube(query):
    pywhatkit.playonyt(query)
    speak(f"Playing {query} on YouTube")


def play_random_song():
    random_searches = [
        "latest Hindi songs",
        "latest English songs",
        "trending songs",
        "Bollywood hits",
        "Top 50 Global",
        "Arijit Singh songs",
        "KK hit songs",
        "Atif Aslam songs",
        "Imagine Dragons",
        "Alan Walker",
        "Taylor Swift songs",
        "Ed Sheeran songs",
        "Lofi music",
        "Punjabi songs",
        "90s Bollywood hits"
    ]

    query = random.choice(random_searches)

    speak("Playing a random song on YouTube.")
    pywhatkit.playonyt(query)


def send_email(recipient, subject, body):
    if not EMAIL_ADDRESS or not EMAIL_PASSWORD:
        speak("Email is not configured. Set OMIRA_EMAIL_ADDRESS and OMIRA_EMAIL_PASSWORD first.")
        return

    email_message = EmailMessage()
    email_message["From"] = EMAIL_ADDRESS
    email_message["To"] = recipient
    email_message["Subject"] = subject
    email_message.set_content(body)

    with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as smtp:
        smtp.starttls()
        smtp.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
        smtp.send_message(email_message)

    speak("Email sent")


# ---------------------------------------------------------------------------
# GMAIL — READ AND REPLY (IMAP)
#
# Uses the same OMIRA_EMAIL_ADDRESS / OMIRA_EMAIL_PASSWORD as send_email().
# For Gmail: turn on 2-Step Verification, then create an "App Password" at
# myaccount.google.com/apppasswords and use that (not your normal password).
# Also make sure IMAP is enabled: Gmail Settings -> Forwarding and POP/IMAP.
# ---------------------------------------------------------------------------

def _imap_connect():
    if not EMAIL_ADDRESS or not EMAIL_PASSWORD:
        raise RuntimeError("email is not configured — set OMIRA_EMAIL_ADDRESS and OMIRA_EMAIL_PASSWORD")
    conn = imaplib.IMAP4_SSL(IMAP_SERVER, IMAP_PORT)
    conn.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
    return conn


def _email_body(msg):
    """Best-effort plain-text body extraction from a parsed email.message.Message."""
    if msg.is_multipart():
        for part in msg.walk():
            if part.get_content_type() == "text/plain" and not part.get("Content-Disposition"):
                try:
                    charset = part.get_content_charset() or "utf-8"
                    return part.get_payload(decode=True).decode(charset, errors="ignore")
                except Exception:
                    continue
        return ""
    try:
        charset = msg.get_content_charset() or "utf-8"
        return msg.get_payload(decode=True).decode(charset, errors="ignore")
    except Exception:
        return ""


def check_email(unread_only=True):
    """Speak a quick summary of the latest inbox messages."""
    global LAST_READ_EMAIL
    try:
        conn = _imap_connect()
    except Exception as exc:
        speak(f"I could not connect to email. {exc}")
        return

    try:
        conn.select("INBOX")
        criterion = "UNSEEN" if unread_only else "ALL"
        status, data = conn.search(None, criterion)
        if status != "OK" or not data[0]:
            speak("No new emails." if unread_only else "Your inbox looks empty.")
            return

        ids = data[0].split()
        latest_ids = ids[-5:][::-1]  # newest 5 first
        speak(f"You have {len(ids)} {'unread ' if unread_only else ''}emails. Here are the latest.")

        for eid in latest_ids:
            status, msg_data = conn.fetch(eid, "(RFC822.HEADER)")
            if status != "OK" or not msg_data or not msg_data[0]:
                continue
            msg = message_from_bytes(msg_data[0][1])
            sender_name, sender_addr = parseaddr(msg.get("From", ""))
            subject = msg.get("Subject", "No subject")
            speak(f"From {sender_name or sender_addr}: {subject}")
            if eid == latest_ids[0]:
                LAST_READ_EMAIL = {"sender": msg.get("From", ""), "subject": subject}
    finally:
        conn.logout()


def read_email_from(sender_query):
    """Read and one-line-summarize the most recent email from a given sender."""
    global LAST_READ_EMAIL
    try:
        conn = _imap_connect()
    except Exception as exc:
        speak(f"I could not connect to email. {exc}")
        return

    try:
        conn.select("INBOX")
        status, data = conn.search(None, f'(FROM "{sender_query}")')
        if status != "OK" or not data[0]:
            speak(f"I could not find an email from {sender_query}.")
            return

        eid = data[0].split()[-1]
        status, msg_data = conn.fetch(eid, "(RFC822)")
        if status != "OK" or not msg_data or not msg_data[0]:
            speak(f"I could not open that email from {sender_query}.")
            return

        msg = message_from_bytes(msg_data[0][1])
        subject = msg.get("Subject", "No subject")
        sender = msg.get("From", sender_query)
        body = _email_body(msg)

        LAST_READ_EMAIL = {"sender": sender, "subject": subject}
        speak(f"Email from {sender_query}, subject: {subject}.")

        if body.strip():
            summary = get_one_line_answer(f"Summarize this email in one sentence: {body[:1500]}")
            speak(summary)
    finally:
        conn.logout()


def reply_to_last_email(body_text):
    """Reply to whichever email was last read via check_email()/read_email_from().
    Goes through the same confirm/cancel flow as other sends."""
    global PENDING_ACTION

    if not LAST_READ_EMAIL.get("sender"):
        speak("I have not read an email yet, so I do not know who to reply to.")
        return

    _, recipient_addr = parseaddr(LAST_READ_EMAIL["sender"])
    if not recipient_addr:
        speak("I could not figure out the sender's email address.")
        return

    subject = LAST_READ_EMAIL.get("subject") or ""
    if not subject.lower().startswith("re:"):
        subject = "Re: " + subject

    PENDING_ACTION = {"type": "send_email", "to": recipient_addr, "subject": subject, "body": body_text}
    speak(f"I am ready to reply to {recipient_addr}. Say confirm to send or cancel to stop.")


def parse_read_email_command(command):
    lowered = command.lower().strip()
    if lowered in {"check my email", "check email", "check inbox", "check my inbox", "any new emails"}:
        return ("check", True)
    if lowered in {"read my inbox", "check all email", "check all emails", "read all email"}:
        return ("check", False)

    m = re.search(r"read (?:the )?email(?:s)? from (?P<sender>.+)", lowered)
    if m:
        return ("read_from", m.group("sender").strip())

    return None


def parse_reply_command(command):
    m = re.search(r"reply(?: to (?:that|it|the email))?\s+saying\s+(?P<body>.+)", command, flags=re.IGNORECASE)
    if m:
        return m.group("body").strip()
    m = re.search(r"reply\s+(?P<body>.+)", command, flags=re.IGNORECASE)
    if m:
        return m.group("body").strip()
    return None


# ---------------------------------------------------------------------------
# WHATSAPP / INSTAGRAM — SEARCH-AND-SEND (no contacts.json needed)
#
# This drives the actual WhatsApp Web / Instagram web UI: it opens the site,
# clicks the search box, types the person's name, opens the top match, and
# types + sends the message. Coordinates below are given as a FRACTION of
# your screen size so they scale to any resolution, but every UI's exact
# layout still varies a bit. If a click misses, run calibrate_coordinates()
# (bottom of this section) to find the right numbers for your screen and
# adjust the *_POS constants below.
#
# Requirements: browser already logged into WhatsApp Web / Instagram, and
# the browser window maximized before Omira clicks anywhere.
# ---------------------------------------------------------------------------

try:
    import pyperclip
except ImportError:
    pyperclip = None

SCREEN_W, SCREEN_H = pyautogui.size()

WHATSAPP_SEARCH_POS = (int(SCREEN_W * 0.09), int(SCREEN_H * 0.08))
WHATSAPP_MESSAGE_BOX_POS = (int(SCREEN_W * 0.45), int(SCREEN_H * 0.95))

INSTAGRAM_SEARCH_POS = (int(SCREEN_W * 0.20), int(SCREEN_H * 0.10))
INSTAGRAM_FIRST_RESULT_POS = (int(SCREEN_W * 0.20), int(SCREEN_H * 0.18))
INSTAGRAM_MESSAGE_BOX_POS = (int(SCREEN_W * 0.55), int(SCREEN_H * 0.90))


def _type_text_fast(text):
    """Type into whatever currently has keyboard focus. Uses clipboard paste
    when available so Hindi/Marathi/emoji text and speed both work."""
    if pyperclip is not None:
        try:
            pyperclip.copy(text)
            pyautogui.hotkey("ctrl", "v")
            return
        except Exception:
            pass
    pyautogui.write(text, interval=0.02)


def calibrate_coordinates(duration=15):
    """Run this on its own to find the right x, y for the *_POS constants
    above: it prints your mouse position once a second. Hover it over the
    WhatsApp/Instagram search box, first result, and message box in turn.

        python -c "import omira_hands_free as j; j.calibrate_coordinates()"
    """
    # Always prints (not gated by OMIRA_VERBOSE) — you run this tool by
    # hand specifically to watch coordinates on screen.
    _print(f"Reporting mouse position for {duration}s — hover over each UI element...")
    end_time = time.time() + duration
    while time.time() < end_time:
        x, y = pyautogui.position()
        _print(f"x={x}, y={y}  (fraction: {x / SCREEN_W:.3f}, {y / SCREEN_H:.3f})")
        time.sleep(1)


# ---------------------------------------------------------------------------
# ONE-LINE Q&A ("ask me anything" mode)
#
# Flow: user asks a question -> Omira speaks ONE short sentence answer.
# Nothing else — no follow-up prompt, straight back to listening after.
# ---------------------------------------------------------------------------

QUESTION_STARTERS = (
    "what", "who", "whom", "whose", "when", "where", "why", "how",
    "is", "are", "was", "were", "do", "does", "did", "can", "could",
    "will", "would", "should", "explain", "define",
    "tell me about", "tell me more about",
)


def looks_like_question(text):
    """True for anything that should get a spoken AI answer rather than be
    treated as a device command. English is detected by question words;
    Hindi/Marathi have no such word list here, so any Devanagari script or
    romanized Hindi/Marathi function word (kya/hai/kay/ahe/etc., reusing the
    same marker sets detect_reply_language() uses) counts too — none of
    Omira's other commands are written in Hindi/Marathi, so text like this
    reaching this point is effectively always a question, and without this
    check it was silently falling through to the website/app planner
    instead of getting answered (looked like "it just opens a webpage")."""
    lowered = text.lower().strip()
    if lowered.endswith("?"):
        return True
    if lowered.startswith(QUESTION_STARTERS):
        return True
    if re.search(r"[\u0900-\u097F]", text):  # Devanagari script
        return True
    words = set(re.findall(r"[a-zA-Z']+", lowered))
    if words & MARATHI_MARKER_WORDS or words & HINDI_MARKER_WORDS:
        return True
    return False


# ---------------------------------------------------------------------------
# SHARED AI CASCADE — Gemini -> OpenAI -> local Ollama
# Every AI-backed feature (one-line answers, document writing, action
# planning) goes through ask_ai() so there's exactly one place to add a
# provider or debug a failure.
# ---------------------------------------------------------------------------

def _call_gemini_one_model(model, system_prompt, user_prompt, max_tokens, request_timeout=8):
    """Try exactly one Gemini model, one network attempt. Returns
    (status, text_or_none) where status is one of:
    'ok', 'try_next_model', 'quota', 'stop'.

    Kept deliberately fast: a single 8s attempt per model, and anything that
    fails (timeout, 5xx, unknown model) moves straight to the next candidate
    model instead of retrying the same one with a longer timeout. The old
    version retried the same model at 15s then 25s before giving up — up to
    40 seconds of dead air on one bad request — which is what made Omira
    feel slow and, combined with only trying one model, is why some
    questions got "I could not reach an AI model" instead of falling over
    to a working model quickly.
    """
    # AQ.-prefixed keys (Google AI Studio's newer "Auth key" format) fail
    # with a 401 "Expected OAuth 2 access token" error if sent via the old
    # ?key= query parameter — they must go in the X-goog-api-key header
    # instead. Sending the header covers both key formats, so this works
    # whether GEMINI_API_KEY is an old AIzaSy... key or a new AQ. one.
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent"
    headers = {"Content-Type": "application/json", "X-goog-api-key": GEMINI_API_KEY}
    REQUEST_TIMEOUT = request_timeout
    global _LAST_AI_ERROR

    def make_payload(output_tokens):
        return {
            "contents": [{"role": "user", "parts": [{"text": user_prompt}]}],
            "systemInstruction": {"parts": [{"text": system_prompt}]},
            "generationConfig": {
                "maxOutputTokens": output_tokens,
                # Keep normal Q&A fast. If a model rejects this field, the
                # request is retried once without it below.
                "thinkingConfig": {"thinkingBudget": 0},
            },
        }

    # Normal attempt, plus (at most) one same-model retry ONLY when Gemini
    # actually responded but got cut off by the output-token budget — that's
    # a cheap, already-fast round trip, not a network-timeout retry.
    output_budgets = [max(128, int(max_tokens))]

    for loop_iteration in range(2):
        payload = make_payload(output_budgets[-1])
        try:
            response = requests.post(url, json=payload, headers=headers, timeout=REQUEST_TIMEOUT)

            # Handle HTTP errors before raise_for_status so we can distinguish
            # quota, authentication, bad model, and temporary server failures.
            if response.status_code >= 400:
                try:
                    error_data = response.json()
                except ValueError:
                    error_data = {}
                message = error_data.get("error", {}).get("message", response.text[:300])

                if response.status_code == 429:
                    log(f"[Gemini] {model}: quota/rate limit (429): {message}")
                    _LAST_AI_ERROR = "your Gemini quota is exhausted for now"
                    return "quota", None

                if response.status_code in (401, 403):
                    log(f"[Gemini] API key/permission error ({response.status_code}): {message}")
                    _LAST_AI_ERROR = "your Gemini API key is invalid, expired, or revoked — generate a new one and update it in .env"
                    return "stop", None

                if response.status_code == 404:
                    log(f"[Gemini] {model}: model not found (404), trying next candidate. {message}")
                    _LAST_AI_ERROR = f"the model '{model}' is not available on your API key"
                    return "try_next_model", None

                if response.status_code in (400, 413):
                    log(f"[Gemini] {model}: bad request ({response.status_code}): {message}")
                    _LAST_AI_ERROR = f"Gemini rejected the request — {message[:150]}"
                    return "stop", None

                if response.status_code >= 500:
                    log(f"[Gemini] {model}: server error ({response.status_code}), trying next candidate.")
                    _LAST_AI_ERROR = "Gemini's servers returned an error"
                    return "try_next_model", None

                log(f"[Gemini] {model}: request failed ({response.status_code}): {message}")
                _LAST_AI_ERROR = f"Gemini returned an unexpected error ({response.status_code})"
                return "try_next_model", None

            data = response.json()
            candidates = data.get("candidates") or []
            if not candidates:
                feedback = data.get("promptFeedback", {})
                log(f"[Gemini] {model}: no candidates. promptFeedback={feedback}")
                _LAST_AI_ERROR = "Gemini returned no answer (the response may have been blocked)"
                return "try_next_model", None

            candidate = candidates[0]
            finish_reason = candidate.get("finishReason", "UNKNOWN")
            parts = candidate.get("content", {}).get("parts") or []
            text = "".join(
                part.get("text", "") for part in parts
                if isinstance(part, dict) and part.get("text")
            ).strip()

            if text:
                _LAST_AI_ERROR = ""
                return "ok", text

            usage = data.get("usageMetadata", {})
            log(f"[Gemini] {model}: empty answer. finishReason={finish_reason}, usage={usage}")

            # Gemini consumed the output budget before emitting text — retry
            # ONCE with a bigger budget (same model, no extra timeout risk).
            if finish_reason == "MAX_TOKENS" and output_budgets[-1] < 1024:
                output_budgets.append(1024)
                log(f"[Gemini] {model}: output limit reached; retrying once with 1024 output tokens...")
                continue

            return "try_next_model", None

        except requests.exceptions.Timeout:
            log(f"[Gemini] {model}: timeout after {REQUEST_TIMEOUT}s, trying next candidate.")
            _LAST_AI_ERROR = "the request to Gemini timed out"
            return "try_next_model", None
        except requests.exceptions.RequestException as exc:
            log(f"[Gemini] {model}: network error: {exc}")
            _LAST_AI_ERROR = "a network error occurred while reaching Gemini"
            return "try_next_model", None
        except (ValueError, TypeError, KeyError) as exc:
            log(f"[Gemini] {model}: invalid response: {exc}")
            _LAST_AI_ERROR = "Gemini returned a response I could not parse"
            return "stop", None
        except Exception as exc:
            log(f"[Gemini] {model}: unexpected error: {exc}")
            _LAST_AI_ERROR = f"an unexpected error occurred: {exc}"
            return "stop", None

    return "try_next_model", None


def _call_gemini(system_prompt, user_prompt, max_tokens, fast=False, deadline=None):
    """Try Gemini across the candidate model list. A model that 404s (been
    renamed/retired) or hits its own quota is skipped in favor of the next
    one — free-tier rate limits are tracked per model, so a different model
    often still has headroom. Whichever model answers first is cached and
    tried first on every later call, so this doesn't add latency once a
    working model is found.

    fast=True (used for one-line Q&A) shortens the per-request timeout and
    caps how many candidate models get tried, so a bad/overloaded model
    can't chew through the whole ~5-7s budget before falling over to
    OpenAI/Ollama.
    """
    global _gemini_cooldown_until, _gemini_working_model, _LAST_AI_ERROR

    if not GEMINI_API_KEY:
        return None

    if time.time() < _gemini_cooldown_until:
        remaining = int(_gemini_cooldown_until - time.time())
        log(f"[Gemini] Skipping call, cooling down for {remaining}s after every candidate model hit quota.")
        _LAST_AI_ERROR = f"your Gemini quota is exhausted — retrying automatically in {remaining}s"
        return None

    # Try the last model that worked first, then the rest of the list.
    order = list(GEMINI_MODEL_CANDIDATES)
    if _gemini_working_model and _gemini_working_model in order:
        order.remove(_gemini_working_model)
        order.insert(0, _gemini_working_model)

    request_timeout = 4 if fast else 8
    if fast:
        order = order[:2]  # cached model + one fallback, never the full list

    any_quota_hit = False
    for model in order:
        if deadline and time.time() > deadline:
            log("[Gemini] Fast-path deadline reached; moving to next backend.")
            _LAST_AI_ERROR = _LAST_AI_ERROR or "Gemini took too long to respond"
            break
        status, text = _call_gemini_one_model(model, system_prompt, user_prompt, max_tokens, request_timeout)
        if status == "ok":
            _gemini_working_model = model
            return text
        if status == "quota":
            any_quota_hit = True
            continue
        if status == "try_next_model":
            continue
        if status == "stop":
            return None

    if any_quota_hit:
        _gemini_cooldown_until = time.time() + GEMINI_COOLDOWN_SECONDS
        log(f"[Gemini] All candidate models are over quota; pausing Gemini for {GEMINI_COOLDOWN_SECONDS}s.")

    return None


def _call_openai(system_prompt, user_prompt, max_tokens, fast=False, deadline=None):
    if not openai_client:
        return None
    try:
        log("[AI] Trying OpenAI fallback...")
        completion = openai_client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            max_tokens=max_tokens,
            timeout=4 if fast else 10,
        )
        content = completion.choices[0].message.content
        return content.strip() if content else None
    except Exception as exc:
        log(f"[OpenAI] Request failed: {exc}")
        return None


def _call_ollama(system_prompt, user_prompt, max_tokens, fast=False, deadline=None):
    try:
        log("[AI] Trying local Ollama fallback...")
        response = requests.post(
            OLLAMA_URL,
            json={
                "model": OLLAMA_MODEL,
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                "stream": False,
                "options": {"num_predict": max_tokens},
            },
            # Local inference gets a bit more slack than the cloud APIs even
            # in fast mode, since there's no network round trip to a remote
            # data center — just the model's own compute time.
            timeout=6 if fast else 20,
        )
        response.raise_for_status()
        content = response.json().get("message", {}).get("content", "")
        return content.strip() if content else None
    except Exception as exc:
        log(f"[Ollama] Request failed: {exc}")
        return None


def ai_backend_configured():
    return bool(GEMINI_API_KEY or openai_client)


def ask_ai(system_prompt, user_prompt, max_tokens=200, fast=False):
    """Try Gemini -> OpenAI -> Ollama without getting stuck on one provider.

    Each backend gets one bounded attempt/retry policy of its own. A provider
    returning no answer immediately moves the request to the next backend.

    fast=True is for latency-sensitive paths (one-line spoken answers): it
    shortens every backend's timeout and stops trying further backends once
    an overall ~7s wall-clock deadline has passed, so Omira never goes
    silent for tens of seconds waiting on a struggling provider.
    """
    backends = (
        ("Gemini", _call_gemini),
        ("OpenAI", _call_openai),
        ("Ollama", _call_ollama),
    )

    global _LAST_AI_ERROR
    _LAST_AI_ERROR = ""  # cleared each call so a stale reason never lingers

    deadline = (time.time() + 7) if fast else None

    for name, call in backends:
        if deadline and time.time() > deadline:
            log(f"[AI] Fast-path deadline reached before trying {name}; giving up.")
            break
        try:
            result = call(system_prompt, user_prompt, max_tokens, fast=fast, deadline=deadline)
            if result:
                log(f"[AI] {name} answered successfully.")
                return result
            log(f"[AI] {name} did not return an answer; moving to the next backend.")
        except Exception as exc:
            # A provider must never break the entire Omira question path.
            log(f"[AI] {name} crashed: {exc}; moving to the next backend.")

    log("[AI] All configured backends failed.")
    return None


def get_one_line_answer(query, language="english"):
    """Real answer to any question, via whichever AI backend is configured,
    in the requested language."""
    system_prompt = (
        OMIRA_PERSONA + f" Answer the user's question in exactly ONE sentence, "
        f"25 to 30 words, written in {LANGUAGE_PROMPT_NAMES.get(language, 'English')}. "
        "No preamble, no lists, no follow-up questions, just the one-sentence answer."
    )
    # 96 output tokens comfortably covers a 30-word sentence in any of these
    # languages while keeping the request small so it comes back fast.
    # fast=True caps the total time spent hunting across backends/models so
    # a one-line answer reliably lands inside ~5-7 seconds instead of
    # potentially chaining through every Gemini candidate model at 8s each.
    answer = ask_ai(system_prompt, query, max_tokens=96, fast=True)
    if answer:
        return answer.splitlines()[0].strip()

    if not ai_backend_configured():
        return (
            "I do not have an AI backend configured, sir — set GEMINI_API_KEY "
            "or OPENAI_API_KEY in your .env, or run Ollama locally."
        )
    if _LAST_AI_ERROR:
        return f"I could not reach an AI model, sir — {_LAST_AI_ERROR}."
    return "I could not reach an AI model just now, sir. Please try again."


def ask_general_question(query):
    """Just the answer — no follow-up prompt, no extra chatter. Answers in
    whichever language the question was asked in."""
    language = detect_reply_language(query)
    speak(get_one_line_answer(query, language=language), language=language)


# ---------------------------------------------------------------------------
# DOCUMENT BUILDER — "write a document about X"
# ---------------------------------------------------------------------------

def _output_folder():
    desktop = Path(os.path.expanduser("~")) / "OneDrive" / "Desktop"
    if not desktop.exists():
        desktop = Path(os.path.expanduser("~")) / "Desktop"
    return desktop


def create_document(topic):
    language = detect_reply_language(topic)
    speak(f"Building a document on {topic}, sir. One moment.", language=language)

    system_prompt = (
        OMIRA_PERSONA + " Write a well-structured, informative document on the "
        f"requested topic, written in {LANGUAGE_PROMPT_NAMES.get(language, 'English')}. "
        "Plain text only, no markdown symbols like # or **. "
        "First line: just the title. Then 3-5 sections, each starting with a "
        "short heading on its own line, followed by 2-4 sentences of clear, "
        "factual prose."
    )
    content = ask_ai(system_prompt, f"Topic: {topic}", max_tokens=1500)

    if not content:
        if not ai_backend_configured():
            speak("I do not have an AI backend configured, sir, so I cannot write that.", language=language)
        else:
            speak("I could not reach an AI model to write that document, sir.", language=language)
        return

    safe_name = re.sub(r"[^a-zA-Z0-9 _-]", "", topic).strip() or "Document"
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    folder = _output_folder()
    lines = [line.strip() for line in content.splitlines() if line.strip()]

    if Document is not None:
        doc = Document()
        if lines:
            doc.add_heading(lines[0], level=0)
            lines = lines[1:]
        for line in lines:
            # Short line with no closing punctuation reads as a heading.
            if len(line) < 60 and not line.endswith((".", ",", ":", ";")):
                doc.add_heading(line, level=1)
            else:
                doc.add_paragraph(line)
        path = folder / f"{safe_name}_{timestamp}.docx"
        doc.save(str(path))
    else:
        path = folder / f"{safe_name}_{timestamp}.txt"
        path.write_text(content, encoding="utf-8")

    try:
        os.startfile(str(path))
    except Exception:
        pass

    speak(f"Document on {topic} is ready, sir.", language=language)


def parse_summarize_command(command):
    m = re.search(r"summar(?:y|ize)\s+(?:of\s+)?(?:the\s+)?(?P<name>.+)", command, flags=re.IGNORECASE)
    return m.group("name").strip() if m else None


def parse_document_command(command):
    patterns = (
        r"(?:write|create|make|build|prepare|generate)\s+(?:me\s+)?(?:a|the)?\s*"
        r"(?:document|report|doc|essay|write[- ]?up)\s+(?:on|about)\s+(?P<topic>.+)",
        r"(?:document|report)\s+(?:on|about)\s+(?P<topic>.+)",
    )
    for pattern in patterns:
        match = re.search(pattern, command, flags=re.IGNORECASE)
        if match:
            return match.group("topic").strip()
    return None


def press_virtual_key(vk_code):
    ctypes.windll.user32.keybd_event(vk_code, 0, 0, 0)
    ctypes.windll.user32.keybd_event(vk_code, 0, 2, 0)


def set_volume_percent(percent):
    if not PYCAW_AVAILABLE:
        speak("Setting exact volume needs the pycaw package installed, sir — nudging it instead.")
        adjust_volume("up" if percent >= 50 else "down")
        return
    try:
        from ctypes import POINTER, cast
        devices = AudioUtilities.GetSpeakers()
        interface = devices.Activate(IAudioEndpointVolume._iid_, CLSCTX_ALL, None)
        volume = cast(interface, POINTER(IAudioEndpointVolume))
        value = max(0, min(100, percent))
        volume.SetMasterVolumeLevelScalar(value / 100, None)
        speak(f"Volume set to {value} percent")
    except Exception as exc:
        speak(f"I could not set the volume, sir. {exc}")


def adjust_volume(direction, steps=6):
    # A single keypress only moves system volume ~2%, which reads as "not
    # working" to a person expecting an audible change. Repeat the press so
    # one voice command produces one clearly noticeable step.
    if direction == "up":
        for _ in range(steps):
            press_virtual_key(VK_VOLUME_UP)
        speak("Volume up")
    elif direction == "down":
        for _ in range(steps):
            press_virtual_key(VK_VOLUME_DOWN)
        speak("Volume down")
    elif direction == "mute":
        press_virtual_key(VK_VOLUME_MUTE)
        speak("Muted")


def change_brightness(delta=None, set_percent=None):
    if sbc is None:
        speak("Brightness control needs the screen-brightness-control package installed, sir.")
        return
    try:
        if set_percent is not None:
            value = max(0, min(100, set_percent))
            sbc.set_brightness(value)
            speak(f"Brightness set to {value} percent")
            return
        current = sbc.get_brightness(display=0)[0]
        value = max(0, min(100, current + delta))
        sbc.set_brightness(value)
        speak(f"Brightness {'increased' if delta > 0 else 'decreased'} to {value} percent")
    except Exception as exc:
        speak(f"I could not change the brightness, sir. {exc}")


def request_power_action(action, label):
    """Shutdown/restart go through the confirm/cancel flow since they're
    disruptive if triggered by a misheard command."""
    global PENDING_ACTION
    PENDING_ACTION = {"type": "power", "action": action}
    speak(f"Ready to {label} the laptop, sir. Say confirm to proceed or cancel to stop.")


def sleep_computer():
    os.system("rundll32.exe powrprof.dll,SetSuspendState 0,1,0")
    speak("Going to sleep, sir.")


def toggle_play_pause():
    press_virtual_key(VK_MEDIA_PLAY_PAUSE)
    speak("Toggled play pause")


def take_screenshot():
    desktop = Path(os.path.expanduser("~")) / "OneDrive" / "Desktop"
    if not desktop.exists():
        desktop = Path(os.path.expanduser("~")) / "Desktop"

    screenshot_path = desktop / f"omira_screenshot_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
    image = pyautogui.screenshot()
    image.save(screenshot_path)
    speak(f"Screenshot saved to {screenshot_path.name}")


def type_text(text):
    pyautogui.write(text, interval=0.01)
    speak("Typed the text")


def press_key(key_name):
    pyautogui.press(key_name)
    speak(f"Pressed {key_name}")


def click_at(x, y):
    pyautogui.click(int(x), int(y))
    speak(f"Clicked at {x} {y}")


def lock_computer():
    os.system("rundll32.exe user32.dll,LockWorkStation")
    speak("Computer locked")


# ---------------------------------------------------------------------------
# EXTRA LAPTOP CONTROL — wifi/bluetooth toggles, recycle bin, quick folders,
# theme, and killing a runaway app by name.
# ---------------------------------------------------------------------------

def set_wifi(enabled):
    """Toggles the first Wi-Fi adapter Windows reports via netsh. Needs an
    elevated (Run as administrator) terminal to actually take effect."""
    try:
        subprocess.run(
            ["netsh", "interface", "set", "interface", "Wi-Fi",
             "enabled" if enabled else "disabled"],
            capture_output=True, timeout=8, check=False,
        )
        speak(f"Wi-Fi turned {'on' if enabled else 'off'}, sir.")
    except Exception as exc:
        speak(f"I could not change Wi-Fi, sir — {exc}")


def set_bluetooth(enabled):
    """Uses the Windows Settings deep link, since toggling the radio itself
    needs a signed driver call that varies by hardware — this is the
    reliable cross-machine way to get there in one command."""
    os.system(f'start ms-settings:bluetooth')
    speak(
        f"Opening Bluetooth settings, sir — toggle it {'on' if enabled else 'off'} there; "
        "every laptop's Bluetooth radio is a little different under the hood."
    )


def empty_recycle_bin():
    try:
        ctypes.windll.shell32.SHEmptyRecycleBinW(None, None, 0x00000001 | 0x00000002)
        speak("Recycle bin emptied, sir.")
    except Exception as exc:
        speak(f"I could not empty the recycle bin, sir — {exc}")


QUICK_FOLDERS = {
    "downloads": "Downloads",
    "documents": "Documents",
    "desktop": "Desktop",
    "pictures": "Pictures",
    "videos": "Videos",
    "music": "Music",
}


def open_quick_folder(name):
    folder_name = QUICK_FOLDERS.get(name.lower().strip())
    if not folder_name:
        speak(f"I don't have a shortcut for the {name} folder, sir.")
        return
    path = Path(os.path.expanduser("~")) / folder_name
    if not path.exists():
        onedrive_path = Path(os.path.expanduser("~")) / "OneDrive" / folder_name
        path = onedrive_path if onedrive_path.exists() else path
    subprocess.Popen(["explorer.exe", str(path)])
    speak(f"Opening {folder_name}, sir.")


def set_dark_mode(dark):
    """Flips the two registry keys Windows checks for app + system theme,
    then nudges Explorer so the change shows immediately."""
    try:
        value = "0" if dark else "1"
        for key_name in ("AppsUseLightTheme", "SystemUsesLightTheme"):
            subprocess.run(
                ["reg", "add",
                 r"HKCU\Software\Microsoft\Windows\CurrentVersion\Themes\Personalize",
                 "/v", key_name, "/t", "REG_DWORD", "/d", value, "/f"],
                capture_output=True, timeout=8, check=False,
            )
        speak(f"Switched to {'dark' if dark else 'light'} mode, sir.")
    except Exception as exc:
        speak(f"I could not change the theme, sir — {exc}")


def close_app(app_name):
    """Force-closes a running app by process/window name, e.g. 'close
    spotify' or 'close chrome'. Best-effort match on the exe name."""
    process_name = app_name.strip().lower().replace(" ", "")
    if not process_name.endswith(".exe"):
        process_name += ".exe"
    result = subprocess.run(
        ["taskkill", "/IM", process_name, "/F"],
        capture_output=True, text=True, timeout=8, check=False,
    )
    if result.returncode == 0:
        speak(f"Closed {app_name}, sir.")
    else:
        speak(f"I could not find {app_name} running, sir.")


def get_live_usage():
    """On-demand CPU/RAM snapshot — separate from get_system_info(), which
    reports static specs. This is the 'how hard is it working right now'
    check, read straight from Windows performance counters."""
    try:
        cpu_output = subprocess.run(
            ["wmic", "cpu", "get", "loadpercentage"],
            capture_output=True, text=True, timeout=6, check=False,
        ).stdout
        cpu_percent = next(
            (line.strip() for line in cpu_output.splitlines()
             if line.strip().isdigit()), "unknown"
        )
    except Exception:
        cpu_percent = "unknown"

    ram_percent = "unknown"
    try:
        class MEMORYSTATUSEX(ctypes.Structure):
            _fields_ = [
                ("dwLength", ctypes.c_ulong), ("dwMemoryLoad", ctypes.c_ulong),
                ("ullTotalPhys", ctypes.c_ulonglong), ("ullAvailPhys", ctypes.c_ulonglong),
                ("ullTotalPageFile", ctypes.c_ulonglong), ("ullAvailPageFile", ctypes.c_ulonglong),
                ("ullTotalVirtual", ctypes.c_ulonglong), ("ullAvailVirtual", ctypes.c_ulonglong),
                ("sullAvailExtendedVirtual", ctypes.c_ulonglong),
            ]
        mem_status = MEMORYSTATUSEX()
        mem_status.dwLength = ctypes.sizeof(MEMORYSTATUSEX)
        if ctypes.windll.kernel32.GlobalMemoryStatusEx(ctypes.byref(mem_status)):
            ram_percent = str(mem_status.dwMemoryLoad)
    except Exception:
        pass

    speak(f"CPU is at {cpu_percent} percent and RAM is at {ram_percent} percent right now, sir.")


def parse_target_payload(target):
    if isinstance(target, dict):
        return target

    if not target:
        return {}

    try:
        return json.loads(target)
    except Exception:
        return {"value": target}


def strip_json_fences(text):
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = cleaned.split("\n", 1)[1] if "\n" in cleaned else ""
        if cleaned.endswith("```"):
            cleaned = cleaned[:-3]
    return cleaned.strip()


def plan_with_local_ai(command):
    system_prompt = (
        OMIRA_PERSONA + " "
        "You control a Windows laptop through a small set of safe tools. "
        "Return only valid JSON with keys action, target, and message. "
        "Allowed actions are open_website, open_app, open_file, search_web, play_song, play_youtube, send_whatsapp, send_email, say, and unknown. "
        "Also allow volume_up, volume_down, mute, play_pause, screenshot, type_text, press_key, click, and lock_computer. "
        "Also allow wifi_on, wifi_off, bluetooth_on, bluetooth_off, empty_recycle_bin, dark_mode, light_mode, live_usage, and close_app. "
        "For send_whatsapp, put JSON in target with phone and message. "
        "For send_email, put JSON in target with to, subject, and body. "
        "For click, put JSON in target with x and y. "
        "For close_app, put the app/process name (e.g. 'spotify') in target. "
        "Use say for normal answers. Use unknown for unsafe or unsupported requests."
    )

    user_prompt = (
        f"User command: {command}\n"
        "Return JSON only, for example: "
        '{"action":"open_app","target":"calculator","message":"Opening Calculator"}'
    )

    content = ask_ai(system_prompt, user_prompt, max_tokens=500)
    if content:
        try:
            return json.loads(strip_json_fences(content))
        except Exception as exc:
            log(f"Could not parse AI plan as JSON: {exc}")

    return {
        "action": "unknown",
        "target": "",
        "message": "I could not reach an AI model just now, sir.",
    }



# ---------------------------------------------------------------------------
# FAST COMMAND ROUTER
# Common commands are handled locally without calling Ollama.
# ---------------------------------------------------------------------------

WEBSITE_ALIASES = {
    "chrome": "https://www.google.com",
    "google": "https://www.google.com",
    "youtube": "https://www.youtube.com",
    "linkedin": "https://www.linkedin.com",
    "instagram": "https://www.instagram.com",
    "whatsapp": "https://web.whatsapp.com",
    "gmail": "https://mail.google.com",
    "mail": "https://mail.google.com",
    "email": "https://mail.google.com",
}

APP_ALIASES = {
    "calculator": "calculator",
    "calc": "calculator",
    "notepad": "notepad",
    "paint": "paint",
    "command prompt": "command prompt",
    "cmd": "command prompt",
    "terminal": "terminal",
    "powershell": "terminal",
    "file explorer": "file explorer",
    "explorer": "file explorer",
    "settings": "settings",
    "control panel": "control panel",
    "task manager": "task manager",
    "device manager": "device manager",
    "this pc": "this pc",
    "my computer": "this pc",
    "wifi settings": "wifi settings",
    "wi-fi settings": "wifi settings",
    "wifi": "wifi settings",
    "bluetooth settings": "bluetooth settings",
    "bluetooth": "bluetooth settings",
    "display settings": "display settings",
    "screen settings": "display settings",
    "sound settings": "sound settings",
    "audio settings": "sound settings",
    "battery settings": "battery settings",
    "storage settings": "storage settings",
    "update settings": "update settings",
    "windows update": "update settings",
    "apps settings": "apps settings",
    "installed apps": "apps settings",
    "network settings": "network settings",
    "printer settings": "printer settings",
    "printers": "printer settings",
}


def parse_instagram_message_command(command):
    """Same idea as parse_message_command but for 'on instagram / on insta'."""
    patterns = [
        r"send\s+(?P<message>.+?)\s+message\s+to\s+(?P<person>.+?)\s+(?:on|via)\s+insta(?:gram)?$",
        r"(?:send\s+)?(?:instagram|insta)\s+message\s+to\s+(?P<person>.+?)\s+(?:saying|that says|message)\s+(?P<message>.+)$",
        r"(?:send\s+)?message\s+to\s+(?P<person>.+?)\s+on\s+insta(?:gram)?\s+(?:saying|that says)\s+(?P<message>.+)$",
        r"send\s+(?P<message>.+?)\s+to\s+(?P<person>.+?)\s+on\s+insta(?:gram)?$",
        r"dm\s+(?P<person>.+?)\s+(?:on\s+insta(?:gram)?\s+)?(?:saying|that says)\s+(?P<message>.+)$",
    ]

    for pattern in patterns:
        match = re.search(pattern, command, flags=re.IGNORECASE)
        if match:
            person = match.group("person").strip()
            message = match.groupdict().get("message", "").strip()
            if message.startswith(("'", '"')) and message.endswith(("'", '"')) and len(message) >= 2:
                message = message[1:-1].strip()
            if not message:
                message = "Hi"
            return person, message

    return None


_OPEN_FILLER_WORDS = {
    "my", "the", "please", "of", "on", "laptop", "computer", "pc",
    "app", "application", "website", "site", "tab", "window", "up",
}


def resolve_open_target(raw_target):
    """Match a spoken 'open X' phrase to a website or app alias, tolerating
    extra words: 'open my gmail', 'open the settings of my laptop', etc."""
    target = raw_target.lower().strip()
    if target in WEBSITE_ALIASES:
        return ("website", WEBSITE_ALIASES[target])
    if target in APP_ALIASES:
        return ("app", APP_ALIASES[target])

    cleaned = " ".join(w for w in target.split() if w not in _OPEN_FILLER_WORDS).strip()
    if cleaned in WEBSITE_ALIASES:
        return ("website", WEBSITE_ALIASES[cleaned])
    if cleaned in APP_ALIASES:
        return ("app", APP_ALIASES[cleaned])

    # Last resort: does a known alias appear as a whole word in what was
    # said? Longest keys checked first so "file explorer" wins over "file".
    for key in sorted(WEBSITE_ALIASES, key=len, reverse=True):
        if re.search(rf"\b{re.escape(key)}\b", target):
            return ("website", WEBSITE_ALIASES[key])
    for key in sorted(APP_ALIASES, key=len, reverse=True):
        if re.search(rf"\b{re.escape(key)}\b", target):
            return ("app", APP_ALIASES[key])

    return None


def process_command(command):
    """Main router: PENDING_ACTION confirmation, then fast deterministic
    matches (apps/websites, WhatsApp/Instagram by name, media/system
    controls), then question-answering, then the local-AI fallback for
    anything else."""
    global PENDING_ACTION
    global CURRENT_LANGUAGE

    raw = command.strip()
    lowered = raw.lower()

    # A previous command is waiting for "confirm" / "cancel".
    if PENDING_ACTION:
        if lowered in {"confirm", "yes", "do it", "send"}:
            action_type = PENDING_ACTION.get("type")
            if action_type == "send_email":
                send_email(PENDING_ACTION["to"], PENDING_ACTION["subject"], PENDING_ACTION["body"])
            elif action_type == "send_whatsapp":
                send_whatsapp_message(
                    PENDING_ACTION["phone"],
                    PENDING_ACTION["message"],
                    PENDING_ACTION.get("person"),
                )
            elif action_type == "power":
                power_action = PENDING_ACTION["action"]
                if power_action == "shutdown":
                    speak("Shutting down, sir.")
                    os.system("shutdown /s /t 5")
                elif power_action == "restart":
                    speak("Restarting, sir.")
                    os.system("shutdown /r /t 5")
            PENDING_ACTION = None
            return
        if lowered in {"cancel", "no", "stop"}:
            PENDING_ACTION = None
            speak("Cancelled")
            return
        # Any other command drops the pending action and continues normally.
        PENDING_ACTION = None

    # Language switch: "speak in hindi" / "hindi mein baat karo" / etc.
    # Sets what Omira LISTENS for going forward (Google's STT needs one
    # fixed language per recording, so this can't be automatic per-question
    # the way replies are).
    new_language = parse_language_switch_command(raw)
    if new_language:
        CURRENT_LANGUAGE = new_language
        confirmations = {
            "english": "Now listening in English, sir.",
            "hindi": "अब मैं हिंदी में सुन रहा हूँ।",
            "marathi": "आता मी मराठीत ऐकत आहे.",
        }
        speak(confirmations[new_language], language=new_language)
        return

    # Websites/apps: exact and simple phrases.
    m = re.fullmatch(r"(?:open|launch|start)\s+(.+)", lowered)
    if m:
        resolved = resolve_open_target(m.group(1))
        if resolved:
            kind, value = resolved
            if kind == "website":
                speak_async(f"Opening {m.group(1).strip()}")
                open_website(value)
            else:
                open_app(value)
            return

    if lowered in {"google", "open google"}:
        open_website(WEBSITE_ALIASES["google"])
        return
    if lowered in {"youtube", "open youtube"}:
        open_website(WEBSITE_ALIASES["youtube"])
        return

    # WhatsApp — send by name, resolved against contacts.json. Doesn't send
    # right away: it stages a PENDING_ACTION and asks for "confirm" first,
    # same pattern as send_email, so a misheard command can't fire off a
    # message unintentionally.
    message_match = parse_message_command(raw)
    if message_match:
        person, message_text = message_match
        phone = resolve_contact(person)
        if not phone:
            speak(f"I don't have a WhatsApp contact saved for {person}, sir. Add them to contacts.json first.")
            return
        PENDING_ACTION = {
            "type": "send_whatsapp",
            "phone": phone,
            "person": person,
            "message": message_text,
        }
        speak(f"Ready to send '{message_text}' to {person} on WhatsApp. Say confirm to send or cancel to stop.")
        return

    # Instagram messaging remains removed from this build -- it relied on
    # pixel-coordinate screen clicking that wasn't reliably verified
    # working. Rather than risk a demo hanging or silently failing, it
    # says so plainly instead of attempting anything.
    if parse_instagram_message_command(raw):
        speak("Instagram messaging isn't included in this build, sir.")
        return

    # Common media/system controls.
    volume_percent = None
    m = re.search(r"(?:set\s+)?volume\s+to\s+(\d{1,3})\s*%?", lowered)
    if m:
        volume_percent = max(0, min(100, int(m.group(1))))
    if volume_percent is not None:
        set_volume_percent(volume_percent)
        return
    if lowered in {"volume up", "increase volume", "louder"} or (
        "volume" in lowered and ("up" in lowered or "increase" in lowered or "raise" in lowered)
    ):
        adjust_volume("up")
        return
    if lowered in {"volume down", "decrease volume", "quieter"} or (
        "volume" in lowered and ("down" in lowered or "decrease" in lowered or "lower" in lowered)
    ):
        adjust_volume("down")
        return
    if lowered in {"mute", "mute volume"}:
        adjust_volume("mute")
        return

    brightness_percent = None
    m = re.search(r"(?:set\s+)?(?:screen\s+)?brightness\s+to\s+(\d{1,3})\s*%?", lowered)
    if m:
        brightness_percent = max(0, min(100, int(m.group(1))))
    if brightness_percent is not None:
        change_brightness(set_percent=brightness_percent)
        return
    if "bright" in lowered and (
        "increase" in lowered or "raise" in lowered or "up" in lowered or lowered == "brighter"
    ):
        change_brightness(delta=15)
        return
    if "bright" in lowered and (
        "decrease" in lowered or "reduce" in lowered or "down" in lowered or "dim" in lowered
    ):
        change_brightness(delta=-15)
        return

    if lowered in {"play pause", "pause", "resume"}:
        toggle_play_pause()
        return
    if lowered in {"screenshot", "take a screenshot"}:
        take_screenshot()
        return
    if lowered in {"lock", "lock computer", "lock my laptop"}:
        lock_computer()
        return
    if lowered in {"shut down my laptop", "shutdown my laptop", "shut down the laptop", "power off"}:
        request_power_action("shutdown", "shut down")
        return
    if lowered in {"restart my laptop", "restart the laptop", "reboot", "reboot my laptop"}:
        request_power_action("restart", "restart")
        return
    if lowered in {"sleep", "sleep my laptop", "put my laptop to sleep", "go to sleep"}:
        sleep_computer()
        return
    if lowered in {"test voice", "voice test", "speak test"}:
        test_voice_output()
        return

    # Extra laptop control — wifi/bluetooth, recycle bin, quick folders,
    # theme, live CPU/RAM, and force-closing a stuck app.
    if lowered in {"wifi on", "turn on wifi", "turn wifi on", "enable wifi"}:
        set_wifi(True)
        return
    if lowered in {"wifi off", "turn off wifi", "turn wifi off", "disable wifi"}:
        set_wifi(False)
        return
    if lowered in {"bluetooth on", "turn on bluetooth", "turn bluetooth on", "enable bluetooth"}:
        set_bluetooth(True)
        return
    if lowered in {"bluetooth off", "turn off bluetooth", "turn bluetooth off", "disable bluetooth"}:
        set_bluetooth(False)
        return
    if lowered in {"empty recycle bin", "empty the recycle bin", "clear recycle bin"}:
        empty_recycle_bin()
        return
    m = re.fullmatch(r"open\s+(downloads|documents|desktop|pictures|videos|music)(?:\s+folder)?", lowered)
    if m:
        open_quick_folder(m.group(1))
        return
    if lowered in {"dark mode", "turn on dark mode", "enable dark mode", "switch to dark mode"}:
        set_dark_mode(True)
        return
    if lowered in {"light mode", "turn on light mode", "enable light mode", "switch to light mode"}:
        set_dark_mode(False)
        return
    if lowered in {"cpu usage", "ram usage", "memory usage", "how hard is my laptop working", "system usage"}:
        get_live_usage()
        return
    m = re.fullmatch(r"close\s+(.+)", lowered)
    if m and m.group(1) not in APP_ALIASES:
        close_app(m.group(1))
        return
    if ("laptop" in lowered or "system" in lowered or "computer" in lowered) and (
        "info" in lowered or "spec" in lowered or "detail" in lowered
    ):
        speak(get_system_info())
        return

    # Battery — "what is my battery percent", "battery status", "how much
    # battery do I have left", "check battery", etc. Excludes "battery
    # settings", which is handled by APP_COMMANDS above.
    if "battery" in lowered and "settings" not in lowered:
        speak(get_battery_status())
        return

    # Document builder — "write a document about X" / "create a report on X".
    doc_topic = parse_document_command(raw)
    if doc_topic:
        create_document(doc_topic)
        return

    # Document summarizer — "summarize the budget report".
    summarize_target = parse_summarize_command(raw)
    if summarize_target:
        speak(summarize_document(summarize_target, ask_ai, OMIRA_PERSONA))
        return

    # Website generator — "build a website for a bakery".
    website_target = parse_website_command(raw)
    if website_target:
        generate_website(website_target, ask_ai, OMIRA_PERSONA, speak_fn=speak)
        return

    # Amazon shopping removed from this build -- not yet reliably
    # verified working end-to-end, so it says so plainly instead of
    # attempting a real purchase flow.
    if lowered.strip() in {"checkout", "check out", "place order"}:
        speak("Shopping isn't included in this build, sir.")
        return
    if re.search(r"^(?:order|buy|shop for)\s+.+", lowered):
        speak("Shopping isn't included in this build, sir.")
        return

    if lowered.startswith("play youtube "):
        play_youtube(command.split(" ", 2)[2])
        return

    if lowered.startswith("play "):
        play_song(command.split(" ", 1)[1])
        return

    # Gmail — read/summarize/reply.
    parsed_check = parse_read_email_command(raw)
    if parsed_check:
        kind, value = parsed_check
        if kind == "check":
            check_email(unread_only=value)
        else:
            read_email_from(value)
        return

    reply_body = parse_reply_command(raw)
    if reply_body:
        reply_to_last_email(reply_body)
        return

    # One-line Q&A mode: "what is X", "how does X work", ends with "?", etc.
    if looks_like_question(raw):
        ask_general_question(raw)
        return

    # Guard against burning AI quota on noise. In hands-free mode (no wake
    # word), the mic hands every stray sound the STT could turn into words —
    # a cough, a TV in the background, one misheard word — to this point.
    # A single short word that isn't a real question or a known short
    # command is almost never a deliberate instruction, so skip the AI
    # planning call for it instead of spending a Gemini/OpenAI request on it.
    if len(lowered.split()) < 2 and not looks_like_question(raw):
        log(f"Ignoring short, ambiguous input (no AI call spent): {raw!r}")
        return

    plan = plan_with_local_ai(command)
    action = plan.get("action", "unknown")
    target = plan.get("target", "")
    message = plan.get("message", "")

    if action == "open_website" and target:
        speak(message or f"Opening {target}")
        open_website(target)
        return

    if action == "open_app" and target:
        open_app(target)
        return

    if action == "open_file" and target:
        open_file(target)
        return

    if action == "search_web" and target:
        speak(message or f"Searching for {target}")
        open_website(f"https://www.google.com/search?q={quote_plus(target)}")
        return

    if action == "play_song" and target:
        play_song(target)
        return

    if action == "play_youtube" and target:
        play_youtube(target)
        return

    if action == "volume_up":
        adjust_volume("up")
        return

    if action == "volume_down":
        adjust_volume("down")
        return

    if action == "mute":
        adjust_volume("mute")
        return

    if action == "play_pause":
        toggle_play_pause()
        return

    if action == "screenshot":
        take_screenshot()
        return

    if action == "type_text" and target:
        type_text(target)
        return

    if action == "press_key" and target:
        press_key(target)
        return

    if action == "click" and target:
        payload = parse_target_payload(target)
        x = payload.get("x")
        y = payload.get("y")
        if x is None or y is None:
            speak("I need x and y coordinates to click.")
            return
        click_at(x, y)
        return

    if action == "lock_computer":
        lock_computer()
        return

    if action == "wifi_on":
        set_wifi(True)
        return

    if action == "wifi_off":
        set_wifi(False)
        return

    if action == "bluetooth_on":
        set_bluetooth(True)
        return

    if action == "bluetooth_off":
        set_bluetooth(False)
        return

    if action == "empty_recycle_bin":
        empty_recycle_bin()
        return

    if action == "dark_mode":
        set_dark_mode(True)
        return

    if action == "light_mode":
        set_dark_mode(False)
        return

    if action == "live_usage":
        get_live_usage()
        return

    if action == "close_app" and target:
        close_app(target)
        return

    if action == "send_whatsapp" and target:
        payload = parse_target_payload(target)
        person_or_phone = payload.get("phone") or payload.get("to") or payload.get("value")
        msg_text = payload.get("message") or message
        if not person_or_phone or not msg_text:
            speak("I need a contact name or phone number and a message to send on WhatsApp.")
            return
        phone = resolve_contact(person_or_phone)
        if not phone:
            speak(f"I don't have a WhatsApp contact saved for {person_or_phone}, sir. Add them to contacts.json first.")
            return
        PENDING_ACTION = {
            "type": "send_whatsapp",
            "phone": phone,
            "person": person_or_phone,
            "message": msg_text,
        }
        speak(f"Ready to send '{msg_text}' to {person_or_phone} on WhatsApp. Say confirm to send or cancel to stop.")
        return

    if action == "send_email" and target:
        payload = parse_target_payload(target)
        recipient = payload.get("to") or payload.get("value")
        subject = payload.get("subject") or "No subject"
        body = payload.get("body") or message
        if not recipient or not body:
            speak("I need a recipient and a body for the email.")
            return
        PENDING_ACTION = {"type": "send_email", "to": recipient, "subject": subject, "body": body}
        speak(f"I am ready to send an email to {recipient}. Say confirm to send or cancel to stop.")
        return

    if action == "say" and message:
        speak(message)
        return

    speak("I'm not sure how to do that yet, sir.")


RUNNING = True  # omira_app.py's tray "Stop" sets this to False to end the loop below


def main():
    global RUNNING, CONTACTS, REMINDERS, RECOGNIZER
    if status_bridge is not None:
        status_bridge.start()

    CONTACTS = load_contacts()
    REMINDERS = load_reminders()

    threading.Thread(target=reminder_worker, daemon=True).start()

    recognizer = create_recognizer()
    RECOGNIZER = recognizer  # share the calibrated recognizer with listen_once()
    initialize_local_stt()

    # Calibrate once at startup. After that Omira continuously listens.
    try:
        with sr.Microphone() as source:
            log("Calibrating microphone...")
            recognizer.adjust_for_ambient_noise(
                source, duration=0.35 if FAST_MODE else 1
            )
    except Exception as exc:
        log(f"Microphone calibration failed: {exc}")
        _status("error", f"microphone calibration failed: {exc}")

    log("Omira hands-free mode is ON.")
    log("Say commands directly — no 'Omira' wake word is required.")
    log("Examples: 'open YouTube', 'send message to Aditya saying hello on WhatsApp'")

    # Always shown (not gated by OMIRA_VERBOSE) — so if Q&A stops working
    # later, you know immediately from the terminal whether it's because no
    # AI backend is configured at all, rather than digging through logs.
    if GEMINI_API_KEY:
        _print(f"AI backend: Gemini ({'/'.join(GEMINI_MODEL_CANDIDATES)})")
    elif openai_client:
        _print(f"AI backend: OpenAI ({OPENAI_MODEL})")
    else:
        _print("AI backend: none configured — set GEMINI_API_KEY or OPENAI_API_KEY in .env, or run Ollama locally.")

    # The only thing spoken at startup — boot diagnostics above stay silent
    # unless OMIRA_VERBOSE=1, but the conversation itself (listening state,
    # what you said, what Omira said) always shows below.
    speak("Hello sir.")

    while RUNNING:
        try:
            _status("listening")
            with sr.Microphone() as source:
                _print("Listening...")
                # timeout: how long to wait for you to start speaking.
                # phrase_time_limit: max length of one recording. Both bumped
                # up slightly so a full spoken question isn't clipped early.
                audio = recognizer.listen(
                    source,
                    timeout=3,
                    phrase_time_limit=10,
                )

            _status("processing", "transcribing")
            command = recognize_audio(audio, recognizer).strip()

            if not command:
                continue

            _print(f"You: {command}")
            _status("processing", command)

            # Execute immediately. No wake word and no "yes, do it" for
            # ordinary commands. Existing send confirmation behavior remains
            # inside the AI fallback for AI-generated email/WhatsApp actions.
            process_command(command)

        except sr.WaitTimeoutError:
            continue
        except sr.UnknownValueError:
            continue
        except sr.RequestError:
            log("Speech recognition service unavailable.")
            _status("error", "speech recognition service unavailable")
            # If using Google STT, retry without killing Omira.
            time.sleep(0.5)
        except KeyboardInterrupt:
            log("\nOmira stopped.")
            _status("idle")
            break
        except Exception as exc:
            log(f"Omira error: {exc}")
            _status("error", str(exc))
            time.sleep(0.2)


if __name__ == "__main__":
    main()